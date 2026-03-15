import os
import base64
import re
import sys
import csv
import io
import json
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from datetime import datetime

import hashlib
import mailchimp_marketing as MailchimpMarketing
from mailchimp_marketing.api_client import ApiClientError

from dotenv import load_dotenv
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from google import genai
from docx import Document

sys.stdout.reconfigure(encoding='utf-8')
load_dotenv()

# --- Config ---
SCOPES = ['https://www.googleapis.com/auth/gmail.readonly']
CREDENTIALS_FILE = 'credentials.json'
TOKEN_FILE = 'token.json'

SENDER_EMAIL = os.getenv('SENDER_EMAIL')
APP_PASSWORD = os.getenv('APP_PASSWORD')
NOTIFY_EMAILS = os.getenv('NOTIFY_EMAILS', '').split(',')
CSV_RECIPIENT = os.getenv('CSV_RECIPIENT')

gemini = genai.Client(api_key=os.getenv('GEMINI_API_KEY'))

MAILCHIMP_API_KEY     = os.getenv('MAILCHIMP_API_KEY')
MAILCHIMP_AUDIENCE_ID = os.getenv('MAILCHIMP_AUDIENCE_ID')

# Derive the data center from the API key (e.g. "us21" from "xxx-us21")
_mc_server = MAILCHIMP_API_KEY.split('-')[-1] if MAILCHIMP_API_KEY else ''
mailchimp = MailchimpMarketing.Client()
mailchimp.set_config({"api_key": MAILCHIMP_API_KEY, "server": _mc_server})

GMAIL_QUERY = (
    f'from:Mercedes@andrewgriffinlawoffice.com '
    f'subject:BK Forms '
    f'has:attachment'
)
LAST_ID_FILE = 'last_processed_id.txt'


# --- Gmail auth ---

def authenticate():
    creds = None
    if os.path.exists(TOKEN_FILE):
        creds = Credentials.from_authorized_user_file(TOKEN_FILE, SCOPES)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(CREDENTIALS_FILE, SCOPES)
            creds = flow.run_local_server(port=0)
        with open(TOKEN_FILE, 'w') as f:
            f.write(creds.to_json())
    return build('gmail', 'v1', credentials=creds)


def get_all_parts(payload):
    """Recursively collect all parts of a Gmail message payload."""
    if 'parts' in payload:
        parts = []
        for part in payload['parts']:
            parts.extend(get_all_parts(part))
        return parts
    return [payload]


# --- Email sending ---

def send_notification(subject, body):
    msg = MIMEText(body)
    msg['Subject'] = subject
    msg['From'] = SENDER_EMAIL
    msg['To'] = ', '.join(NOTIFY_EMAILS)
    _send(NOTIFY_EMAILS, msg)
    print("Notification sent.")


def send_csv_email(csv_data, run_time):
    msg = MIMEMultipart()
    msg['Subject'] = f"[BK Forms] Parsed Client Info - {run_time}"
    msg['From'] = SENDER_EMAIL
    msg['To'] = CSV_RECIPIENT
    msg.attach(MIMEText("Attached is the CSV with client contact info parsed from today's intake forms."))

    attachment = MIMEBase('application', 'octet-stream')
    attachment.set_payload(csv_data.encode('utf-8'))
    encoders.encode_base64(attachment)
    safe_time = run_time.replace(':', '-').replace(' ', '_')
    attachment.add_header('Content-Disposition', f'attachment; filename="client_info_{safe_time}.csv"')
    msg.attach(attachment)

    _send([CSV_RECIPIENT], msg)
    print(f"CSV sent to {CSV_RECIPIENT}.")


def _send(recipients, msg):
    try:
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
            smtp.login(SENDER_EMAIL, APP_PASSWORD)
            smtp.sendmail(SENDER_EMAIL, recipients, msg.as_string())
    except Exception as e:
        print(f"Failed to send email: {e}")


def client_prefix(filename):
    """Extract the client name from 'LastName, First - Form Title.ext' → 'lastname, first'"""
    parts = filename.split(' - ', 1)
    return parts[0].strip().lower()


def download_attachment(service, message_id, part):
    body = part.get('body', {})
    if 'attachmentId' in body:
        result = service.users().messages().attachments().get(
            userId='me', messageId=message_id, id=body['attachmentId']
        ).execute()
        return base64.urlsafe_b64decode(result['data'])
    elif 'data' in body:
        return base64.urlsafe_b64decode(body['data'])
    return None


# --- Gemini parsing ---

def extract_docx_text(file_bytes):
    doc = Document(io.BytesIO(file_bytes))
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text.strip())
    return '\n'.join(lines)


def parse_client_info(docx_text, filename):
    """Ask Gemini to extract first name, last name, email, and phone from form text."""
    prompt = f"""Extract contact information from the legal intake form below.
Return ONLY a JSON object with these keys: first_name, last_name, email, phone.
Use null for any field not found. No extra text or formatting.

{docx_text}"""

    try:
        response = gemini.models.generate_content(model='gemini-2.5-flash', contents=prompt)
        raw = re.sub(r'^```(?:json)?\s*|\s*```$', '', response.text.strip())
        data = json.loads(raw)
        print(f"  Parsed: {data}")
        return data
    except Exception as e:
        print(f"  Gemini parsing failed for {filename}: {e}")
        return None


def mailchimp_subscribe(info):
    """Add contact if new, otherwise leave existing data untouched."""
    email = info.get('email', '').strip().lower()
    if not email:
        print("  Skipping Mailchimp: no email address.")
        return False
    subscriber_hash = hashlib.md5(email.encode()).hexdigest()
    try:
        mailchimp.lists.get_list_member(MAILCHIMP_AUDIENCE_ID, subscriber_hash)
        print(f"  Mailchimp: contact already exists — skipping subscribe ({email})")
        return True
    except ApiClientError:
        pass  # 404 means contact doesn't exist yet — create them below

    try:
        mailchimp.lists.add_list_member(
            MAILCHIMP_AUDIENCE_ID,
            {
                "email_address": email,
                "status": "subscribed",
                "merge_fields": {
                    "FNAME": info.get('first_name') or '',
                    "LNAME": info.get('last_name') or '',
                    "PHONE": info.get('phone') or '',
                },
            },
        )
        print(f"  Mailchimp: new contact subscribed ({email})")
        return True
    except ApiClientError as e:
        if 'resubscribe' in e.text.lower() or 're-subscribe' in e.text.lower():
            print(f"  Mailchimp: contact is unsubscribed — skipping subscribe, will still tag ({email})")
            return True
        print(f"  Mailchimp subscribe failed for {email}: {e.text}")
        return False


def mailchimp_add_tag(info, tag="Bankruptcy Lead"):
    """Apply a tag to a contact, triggering any existing automation."""
    email = info.get('email', '').strip().lower()
    if not email:
        return
    subscriber_hash = hashlib.md5(email.encode()).hexdigest()
    try:
        mailchimp.lists.update_list_member_tags(
            MAILCHIMP_AUDIENCE_ID,
            subscriber_hash,
            {"tags": [{"name": tag, "status": "active"}]},
        )
        print(f"  Mailchimp: tagged '{tag}' → {email}")
    except ApiClientError as e:
        print(f"  Mailchimp tagging failed for {email}: {e.text}")


def build_csv(rows):
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=['first_name', 'last_name', 'email', 'phone'])
    writer.writeheader()
    writer.writerows(rows)
    return output.getvalue()


def load_last_id():
    if os.path.exists(LAST_ID_FILE):
        return open(LAST_ID_FILE).read().strip()
    return None


def save_last_id(msg_id):
    with open(LAST_ID_FILE, 'w') as f:
        f.write(msg_id)


def check_for_new_email(service):
    """Return the latest email if it hasn't been processed yet, otherwise None."""
    last_id = load_last_id()
    results = service.users().messages().list(userId='me', q=GMAIL_QUERY, maxResults=1).execute()
    messages = results.get('messages', [])
    if messages and messages[0]['id'] != last_id:
        return messages[0]
    return None


# --- Main ---

def main():
    run_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    print(f"\n--- Run started: {run_time} ---")

    try:
        service = authenticate()
        print("Authenticated successfully!\n")

        msg_ref = check_for_new_email(service)
        if not msg_ref:
            print("No new email found. Exiting.")
            return
        meta = service.users().messages().get(
            userId='me', id=msg_ref['id'], format='metadata',
            metadataHeaders=['Subject', 'Date', 'From']
        ).execute()
        headers = {h['name']: h['value'] for h in meta['payload']['headers']}
        email_subject = headers.get('Subject', 'No Subject')

        print(f"Subject: {email_subject}")
        print(f"From:    {headers.get('From', 'Unknown')}")
        print(f"Date:    {headers.get('Date', '')}\n")

        # Fetch full message and process attachments
        full_msg = service.users().messages().get(userId='me', id=msg_ref['id']).execute()
        parts = get_all_parts(full_msg.get('payload', {}))

        # Build set of client names that have a paired .pdf
        pdf_clients = {
            client_prefix(part['filename'])
            for part in parts
            if part.get('filename', '').lower().endswith('.pdf')
        }
        if pdf_clients:
            print(f"Paired .pdf clients (Mailchimp only, excluded from CSV): {pdf_clients}\n")

        parsed_rows = []
        paired_count = 0

        for part in parts:
            filename = part.get('filename')
            if not filename:
                continue

            if not filename.lower().endswith('.docx'):
                continue

            is_paired = client_prefix(filename) in pdf_clients
            print(f"Found: {filename}" + (" (paired)" if is_paired else ""))

            data = download_attachment(service, msg_ref['id'], part)
            if data is None:
                print(f"  Could not download: {filename}")
                continue

            print(f"  Downloaded ({len(data):,} bytes), parsing with Gemini...")
            text = extract_docx_text(data)
            info = parse_client_info(text, filename)
            if info:
                if not is_paired:
                    parsed_rows.append({
                        'first_name': info.get('first_name') or '',
                        'last_name':  info.get('last_name') or '',
                        'email':      info.get('email') or '',
                        'phone':      info.get('phone') or '',
                    })
                else:
                    paired_count += 1
                if mailchimp_subscribe(info):
                    mailchimp_add_tag(info)
            print()

        save_last_id(msg_ref['id'])
        print(f"Done! Parsed {len(parsed_rows)} attachment(s).")

        if parsed_rows:
            send_csv_email(build_csv(parsed_rows), run_time)
            send_notification(
                "[BK Forms] Success - Client Info CSV Sent",
                f"Run time: {run_time}\n\nParsed {len(parsed_rows)} .docx attachment(s).\n"
                f"CSV sent to {CSV_RECIPIENT}.\n\nEmail: {email_subject}"
            )
        elif paired_count:
            send_notification(
                "[BK Forms] All Clients Paired - Mailchimp Only",
                f"Run time: {run_time}\n\n{paired_count} client(s) were processed and tagged in Mailchimp.\n"
                f"All had paired .pdf files so no CSV was generated.\n\nEmail: {email_subject}"
            )
        else:
            send_notification(
                "[BK Forms] No .docx Attachments Found",
                f"Run time: {run_time}\n\nNo .docx attachments were found or parsed.\n\nEmail: {email_subject}"
            )

    except Exception as e:
        error_msg = f"Script failed: {e}"
        print(error_msg)
        send_notification("[BK Forms] ERROR - Script Failed", f"Run time: {run_time}\n\n{error_msg}")


if __name__ == '__main__':
    main()
